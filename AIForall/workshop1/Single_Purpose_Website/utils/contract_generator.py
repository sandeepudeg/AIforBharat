from models.contract import ContractData
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib import colors
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime

class ContractGenerator:
    """Generate contracts in PDF and Word formats"""
    
    UPLOAD_FOLDER = os.path.join(os.path.dirname(__file__), '..', 'uploads')
    
    @staticmethod
    def format_contract_html(contract: ContractData) -> str:
        """Format contract as HTML for preview"""
        html = f"""
        <div class="contract-preview">
            <h1>CONTRACT AGREEMENT</h1>
            
            <h2>1. PARTIES INFORMATION</h2>
            <p><strong>Party 1:</strong> {contract.party1_name}</p>
            <p><strong>Address:</strong> {contract.party1_address}</p>
            <p><strong>Entity Type:</strong> {contract.party1_entity_type}</p>
            
            <p><strong>Party 2:</strong> {contract.party2_name}</p>
            <p><strong>Address:</strong> {contract.party2_address}</p>
            <p><strong>Entity Type:</strong> {contract.party2_entity_type}</p>
            
            <h2>2. PURPOSE & SCOPE</h2>
            <p><strong>Purpose:</strong> {contract.contract_purpose}</p>
            <p><strong>Scope of Work:</strong> {contract.scope_of_work}</p>
            <p><strong>Deliverables:</strong> {contract.deliverables}</p>
            
            <h2>3. KEY TERMS</h2>
            <p><strong>Start Date:</strong> {contract.start_date}</p>
            <p><strong>End Date:</strong> {contract.end_date}</p>
            <p><strong>Payment Amount:</strong> {contract.payment_amount}</p>
            <p><strong>Payment Schedule:</strong> {contract.payment_schedule}</p>
            {f'<p><strong>Late Payment Penalty:</strong> {contract.late_payment_penalty}</p>' if contract.late_payment_penalty else ''}
            {f'<p><strong>Performance Standards:</strong> {contract.performance_standards}</p>' if contract.performance_standards else ''}
            
            <h2>4. LEGAL COMPLIANCE</h2>
            {f'<p><strong>Legal Compliance:</strong> {contract.legal_compliance}</p>' if contract.legal_compliance else ''}
            {f'<p><strong>Licenses & Permits:</strong> {contract.licenses_permits}</p>' if contract.licenses_permits else ''}
            
            <h2>5. RISK & LIABILITY</h2>
            {f'<p><strong>Liability Clauses:</strong> {contract.liability_clauses}</p>' if contract.liability_clauses else ''}
            {f'<p><strong>Indemnity Provisions:</strong> {contract.indemnity_provisions}</p>' if contract.indemnity_provisions else ''}
            {f'<p><strong>Insurance Requirements:</strong> {contract.insurance_requirements}</p>' if contract.insurance_requirements else ''}
            
            <h2>6. CONFIDENTIALITY & IP</h2>
            {f'<p><strong>Confidentiality Obligations:</strong> {contract.confidentiality_obligations}</p>' if contract.confidentiality_obligations else ''}
            {f'<p><strong>IP Ownership:</strong> {contract.ip_ownership}</p>' if contract.ip_ownership else ''}
            
            <h2>7. TERMINATION</h2>
            {f'<p><strong>Termination Conditions:</strong> {contract.termination_conditions}</p>' if contract.termination_conditions else ''}
            {f'<p><strong>Notice Period:</strong> {contract.notice_period}</p>' if contract.notice_period else ''}
            {f'<p><strong>Termination Consequences:</strong> {contract.termination_consequences}</p>' if contract.termination_consequences else ''}
            
            <h2>8. DISPUTE RESOLUTION</h2>
            <p><strong>Method:</strong> {contract.dispute_resolution_method}</p>
            <p><strong>Jurisdiction:</strong> {contract.jurisdiction}</p>
            <p><strong>Governing Law:</strong> {contract.governing_law}</p>
            
            <h2>9. SIGNATURES</h2>
            <p>Party 1: _________________________ Date: _________</p>
            <p>{contract.party1_name}</p>
            
            <p>Party 2: _________________________ Date: _________</p>
            <p>{contract.party2_name}</p>
        </div>
        """
        return html
    
    @staticmethod
    def generate_pdf(contract: ContractData) -> str:
        """Generate PDF contract"""
        os.makedirs(ContractGenerator.UPLOAD_FOLDER, exist_ok=True)
        
        filename = f"contract_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        filepath = os.path.join(ContractGenerator.UPLOAD_FOLDER, filename)
        
        # Create PDF
        doc = SimpleDocTemplate(filepath, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            textColor=colors.HexColor('#1a1a1a'),
            spaceAfter=12,
            alignment=1  # Center
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=11,
            textColor=colors.HexColor('#333333'),
            spaceAfter=6,
            spaceBefore=6
        )
        
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=9,
            textColor=colors.HexColor('#333333'),
            spaceAfter=4
        )
        
        # Title
        story.append(Paragraph("CONTRACT AGREEMENT", title_style))
        story.append(Spacer(1, 0.1*inch))
        
        # Parties Information
        story.append(Paragraph("1. PARTIES INFORMATION", heading_style))
        story.append(Paragraph(f"<b>Party 1:</b> {contract.party1_name}", body_style))
        story.append(Paragraph(f"<b>Address:</b> {contract.party1_address}", body_style))
        story.append(Paragraph(f"<b>Entity Type:</b> {contract.party1_entity_type}", body_style))
        story.append(Spacer(1, 0.05*inch))
        story.append(Paragraph(f"<b>Party 2:</b> {contract.party2_name}", body_style))
        story.append(Paragraph(f"<b>Address:</b> {contract.party2_address}", body_style))
        story.append(Paragraph(f"<b>Entity Type:</b> {contract.party2_entity_type}", body_style))
        
        # Purpose & Scope
        story.append(Paragraph("2. PURPOSE & SCOPE", heading_style))
        story.append(Paragraph(f"<b>Purpose:</b> {contract.contract_purpose}", body_style))
        story.append(Paragraph(f"<b>Scope:</b> {contract.scope_of_work}", body_style))
        story.append(Paragraph(f"<b>Deliverables:</b> {contract.deliverables}", body_style))
        
        # Key Terms
        story.append(Paragraph("3. KEY TERMS", heading_style))
        story.append(Paragraph(f"<b>Duration:</b> {contract.start_date} to {contract.end_date}", body_style))
        story.append(Paragraph(f"<b>Payment:</b> {contract.payment_amount} ({contract.payment_schedule})", body_style))
        if contract.late_payment_penalty:
            story.append(Paragraph(f"<b>Late Payment Penalty:</b> {contract.late_payment_penalty}", body_style))
        
        # Legal Compliance
        story.append(Paragraph("4. LEGAL COMPLIANCE", heading_style))
        if contract.legal_compliance:
            story.append(Paragraph(f"{contract.legal_compliance}", body_style))
        if contract.licenses_permits:
            story.append(Paragraph(f"<b>Licenses & Permits:</b> {contract.licenses_permits}", body_style))
        
        # Risk & Liability
        story.append(Paragraph("5. RISK & LIABILITY", heading_style))
        if contract.liability_clauses:
            story.append(Paragraph(f"<b>Liability:</b> {contract.liability_clauses}", body_style))
        if contract.indemnity_provisions:
            story.append(Paragraph(f"<b>Indemnity:</b> {contract.indemnity_provisions}", body_style))
        
        # Confidentiality & IP
        story.append(Paragraph("6. CONFIDENTIALITY & IP", heading_style))
        if contract.confidentiality_obligations:
            story.append(Paragraph(f"<b>Confidentiality:</b> {contract.confidentiality_obligations}", body_style))
        if contract.ip_ownership:
            story.append(Paragraph(f"<b>IP Ownership:</b> {contract.ip_ownership}", body_style))
        
        # Termination
        story.append(Paragraph("7. TERMINATION", heading_style))
        if contract.termination_conditions:
            story.append(Paragraph(f"<b>Conditions:</b> {contract.termination_conditions}", body_style))
        if contract.notice_period:
            story.append(Paragraph(f"<b>Notice Period:</b> {contract.notice_period}", body_style))
        
        # Dispute Resolution
        story.append(Paragraph("8. DISPUTE RESOLUTION", heading_style))
        story.append(Paragraph(f"<b>Method:</b> {contract.dispute_resolution_method}", body_style))
        story.append(Paragraph(f"<b>Jurisdiction:</b> {contract.jurisdiction}", body_style))
        story.append(Paragraph(f"<b>Governing Law:</b> {contract.governing_law}", body_style))
        
        # Signatures
        story.append(Spacer(1, 0.1*inch))
        story.append(Paragraph("9. SIGNATURES", heading_style))
        story.append(Paragraph(f"Party 1: _________________________ Date: _________", body_style))
        story.append(Paragraph(f"{contract.party1_name}", body_style))
        story.append(Spacer(1, 0.05*inch))
        story.append(Paragraph(f"Party 2: _________________________ Date: _________", body_style))
        story.append(Paragraph(f"{contract.party2_name}", body_style))
        
        # Build PDF
        doc.build(story)
        
        return filepath
    
    @staticmethod
    def generate_docx(contract: ContractData) -> str:
        """Generate Word document contract"""
        os.makedirs(ContractGenerator.UPLOAD_FOLDER, exist_ok=True)
        
        filename = f"contract_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(ContractGenerator.UPLOAD_FOLDER, filename)
        
        # Create Word document
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Title
        title = doc.add_paragraph()
        title_run = title.add_run("CONTRACT AGREEMENT")
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Parties Information
        doc.add_paragraph("1. PARTIES INFORMATION", style='Heading 2')
        doc.add_paragraph(f"Party 1: {contract.party1_name}")
        doc.add_paragraph(f"Address: {contract.party1_address}")
        doc.add_paragraph(f"Entity Type: {contract.party1_entity_type}")
        doc.add_paragraph(f"Party 2: {contract.party2_name}")
        doc.add_paragraph(f"Address: {contract.party2_address}")
        doc.add_paragraph(f"Entity Type: {contract.party2_entity_type}")
        
        # Purpose & Scope
        doc.add_paragraph("2. PURPOSE & SCOPE", style='Heading 2')
        doc.add_paragraph(f"Purpose: {contract.contract_purpose}")
        doc.add_paragraph(f"Scope of Work: {contract.scope_of_work}")
        doc.add_paragraph(f"Deliverables: {contract.deliverables}")
        
        # Key Terms
        doc.add_paragraph("3. KEY TERMS", style='Heading 2')
        doc.add_paragraph(f"Duration: {contract.start_date} to {contract.end_date}")
        doc.add_paragraph(f"Payment: {contract.payment_amount} ({contract.payment_schedule})")
        if contract.late_payment_penalty:
            doc.add_paragraph(f"Late Payment Penalty: {contract.late_payment_penalty}")
        if contract.performance_standards:
            doc.add_paragraph(f"Performance Standards: {contract.performance_standards}")
        
        # Legal Compliance
        doc.add_paragraph("4. LEGAL COMPLIANCE", style='Heading 2')
        if contract.legal_compliance:
            doc.add_paragraph(contract.legal_compliance)
        if contract.licenses_permits:
            doc.add_paragraph(f"Licenses & Permits: {contract.licenses_permits}")
        
        # Risk & Liability
        doc.add_paragraph("5. RISK & LIABILITY", style='Heading 2')
        if contract.liability_clauses:
            doc.add_paragraph(f"Liability: {contract.liability_clauses}")
        if contract.indemnity_provisions:
            doc.add_paragraph(f"Indemnity: {contract.indemnity_provisions}")
        if contract.insurance_requirements:
            doc.add_paragraph(f"Insurance: {contract.insurance_requirements}")
        
        # Confidentiality & IP
        doc.add_paragraph("6. CONFIDENTIALITY & IP", style='Heading 2')
        if contract.confidentiality_obligations:
            doc.add_paragraph(f"Confidentiality: {contract.confidentiality_obligations}")
        if contract.ip_ownership:
            doc.add_paragraph(f"IP Ownership: {contract.ip_ownership}")
        
        # Termination
        doc.add_paragraph("7. TERMINATION", style='Heading 2')
        if contract.termination_conditions:
            doc.add_paragraph(f"Conditions: {contract.termination_conditions}")
        if contract.notice_period:
            doc.add_paragraph(f"Notice Period: {contract.notice_period}")
        if contract.termination_consequences:
            doc.add_paragraph(f"Consequences: {contract.termination_consequences}")
        
        # Dispute Resolution
        doc.add_paragraph("8. DISPUTE RESOLUTION", style='Heading 2')
        doc.add_paragraph(f"Method: {contract.dispute_resolution_method}")
        doc.add_paragraph(f"Jurisdiction: {contract.jurisdiction}")
        doc.add_paragraph(f"Governing Law: {contract.governing_law}")
        
        # Signatures
        doc.add_paragraph()
        doc.add_paragraph("9. SIGNATURES", style='Heading 2')
        doc.add_paragraph(f"Party 1: _________________________ Date: _________")
        doc.add_paragraph(f"{contract.party1_name}")
        doc.add_paragraph()
        doc.add_paragraph(f"Party 2: _________________________ Date: _________")
        doc.add_paragraph(f"{contract.party2_name}")
        
        # Save document
        doc.save(filepath)
        
        return filepath
